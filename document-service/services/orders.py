import datetime
import os
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_and_upload_vacation_schedule_order(
    count_order,
    company_name='ООО «ТНС энерго Пенза»',
    hr_manager_name='Солевашин А. А.',
    output_filename='Приказ_график_отпусков.docx',
):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    company_para = doc.add_paragraph(company_name)
    company_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    company_para.runs[0].font.bold = True

    title_para = doc.add_paragraph(f'ПРИКАЗ № {count_order}')
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.runs[0].font.bold = True
    title_para.runs[0].font.size = Pt(14)
    date = datetime.datetime.now()
    date_para = doc.add_paragraph(f"Пенза\t\t\t\t\t\t\t\t{date.strftime('17.12.%Y')}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    text = doc.add_paragraph('Об утверждении графика отпусков')
    text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    text.runs[0].font.bold = True

    doc.add_paragraph()

    doc.add_paragraph('Во исполнение обязанности, предусмотренной ст. 123 ТК РФ:')

    doc.add_paragraph()

    command_para = doc.add_paragraph('ПРИКАЗЫВАЮ:')
    command_para.runs[0].font.bold = True


    y = date.year + 1
    doc.add_paragraph(
        'Утвердить график отпусков работников ' + company_name + f' на {y} год согласно приложению.',
        style='List Number'
    )

    doc.add_paragraph(
        'Начальнику отдела кадров ' + hr_manager_name + ' ознакомить работников ' +
        company_name + ' с утвержденным графиком отпусков под роспись и обеспечить его соблюдение в течение 2026 года.',
        style='List Number'
    )

    doc.add_paragraph('\n\nГенеральный директор')
    doc.add_paragraph('_________________')

    doc.save(output_filename)

    with open(output_filename, 'rb') as f:
        files = {
            'file': (output_filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }

        headers = {
            'accept': 'application/json'
        }

        response = requests.post(
            'https://hackathon-calendar.duckdns.org/file-service/file',
            headers=headers,
            files=files
        )

    os.remove(output_filename)
    return response.json()
